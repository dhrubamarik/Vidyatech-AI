"""
Turns a chat response (which may contain light markdown like **bold**,
bullet points, and headers) into a downloadable PDF or DOCX file.

This is intentionally a simple, readable renderer rather than a full
markdown engine -- agent responses are short-form text, not complex
documents, so a few line-level rules cover the vast majority of output.
"""
import io
import re
from typing import List, Tuple

from docx import Document
from docx.shared import Pt

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.enums import TA_LEFT

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _classify_lines(text: str) -> List[Tuple[str, str]]:
    """Return a list of (kind, content) tuples: kind is 'bullet', 'heading',
    or 'para'."""
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith(("- ", "* ")):
            lines.append(("bullet", line[2:].strip()))
        elif line.startswith("#"):
            lines.append(("heading", line.lstrip("#").strip()))
        else:
            lines.append(("para", line))
    return lines


def generate_docx_bytes(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title or "VidyaTech AI Response", level=1)

    for kind, text in _classify_lines(content):
        clean = BOLD_RE.sub(r"\1", text)  # docx bold handled per-run below
        if kind == "heading":
            doc.add_heading(clean, level=2)
        elif kind == "bullet":
            doc.add_paragraph(clean, style="List Bullet")
        else:
            p = doc.add_paragraph()
            # Preserve **bold** as real bold runs instead of just stripping it
            pos = 0
            for m in BOLD_RE.finditer(text):
                if m.start() > pos:
                    p.add_run(text[pos:m.start()])
                run = p.add_run(m.group(1))
                run.bold = True
                pos = m.end()
            if pos < len(text):
                p.add_run(text[pos:])
            for run in p.runs:
                run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_pdf_bytes(title: str, content: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, topMargin=0.8 * inch, bottomMargin=0.8 * inch)
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=11, leading=16, alignment=TA_LEFT, spaceAfter=8)
    heading_style = ParagraphStyle("Heading", parent=styles["Heading2"], spaceBefore=10, spaceAfter=6)

    story = [Paragraph(title or "VidyaTech AI Response", styles["Title"]), Spacer(1, 14)]

    bullet_buffer = []

    def flush_bullets():
        if bullet_buffer:
            story.append(ListFlowable(
                [ListItem(Paragraph(_to_reportlab_markup(b), body_style)) for b in bullet_buffer],
                bulletType="bullet"
            ))
            bullet_buffer.clear()

    for kind, text in _classify_lines(content):
        if kind == "bullet":
            bullet_buffer.append(text)
            continue
        flush_bullets()
        if kind == "heading":
            story.append(Paragraph(_to_reportlab_markup(text), heading_style))
        else:
            story.append(Paragraph(_to_reportlab_markup(text), body_style))
    flush_bullets()

    doc.build(story)
    buf.seek(0)
    return buf.read()


def _to_reportlab_markup(text: str) -> str:
    """ReportLab's Paragraph understands a small HTML-like subset,
    including <b>. Convert **bold** into that, and escape the few
    characters that would otherwise break its parser."""
    escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return BOLD_RE.sub(r"<b>\1</b>", escaped)
